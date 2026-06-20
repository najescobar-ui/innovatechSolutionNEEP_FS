#!/usr/bin/env python3
"""Genera report.pdf y report.docx a partir de report.md con la identidad Duoc UC.
PDF via reportlab, DOCX via python-docx (sin LibreOffice). Una sola fuente: report.md.
"""
import html
import os
import re

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "assets")
FONTS = os.path.join(ASSETS, "fonts")
LOGO = os.path.join(ASSETS, "duoc_logo.png")
MD = os.path.join(BASE, "report.md")
PDF = os.path.join(BASE, "report.pdf")
DOCX = os.path.join(BASE, "report.docx")

# Paleta Duoc UC
AZUL = "#002138"
DORADO = "#FFB800"
GRIS = "#666666"
NEGRO = "#1A1A1A"
ZEBRA = "#F4F4F4"
CODEBG = "#F5F5F5"

INLINE = re.compile(r"(`[^`]+`)|(\*\*.+?\*\*)|(\[[^\]]+\]\([^)]*\))")
LINKRE = re.compile(r"\[([^\]]+)\]\(([^)]*)\)")


# ---------------------------------------------------------------- parser md
def parse_md(text):
    lines = text.split("\n")
    h1 = None
    meta = {}
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == "":
            i += 1
            continue
        if line.startswith("# ") and h1 is None:
            h1 = line[2:].strip()
            i += 1
            continue
        m = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", s)
        if m and "|" not in s:
            meta[m.group(1)] = m.group(2)
            i += 1
            continue
        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(buf)))
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if s == "---":
            blocks.append(("hr", ""))
            i += 1
            continue
        if line.startswith("> "):
            buf = []
            while i < n and lines[i].startswith("> "):
                buf.append(lines[i][2:])
                i += 1
            blocks.append(("quote", " ".join(buf)))
            continue
        if line.startswith("|"):
            buf = []
            while i < n and lines[i].lstrip().startswith("|"):
                buf.append(lines[i])
                i += 1
            rows = []
            for r in buf:
                if re.match(r"^\s*\|[\s:\-|]+\|?\s*$", r):
                    continue
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append(("table", rows))
            continue
        if line.startswith("- "):
            items = []
            while i < n and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append(("ul", items))
            continue
        # parrafo: juntar lineas consecutivas "normales"
        buf = []
        while i < n:
            ln = lines[i]
            st = ln.strip()
            if st == "" or ln.startswith(("#", "> ", "|", "- ", "```")) or st == "---":
                break
            buf.append(st)
            i += 1
        blocks.append(("p", " ".join(buf)))
    return h1, meta, blocks


# ---------------------------------------------------------------- PDF
def build_pdf(h1, meta, blocks):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, Image, PageBreak, HRFlowable,
                                     Preformatted)
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    pdfmetrics.registerFont(TTFont("Merri", os.path.join(FONTS, "Merriweather-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Merri-B", os.path.join(FONTS, "Merriweather-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("Lato", os.path.join(FONTS, "Lato-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Lato-B", os.path.join(FONTS, "Lato-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("Lato-I", os.path.join(FONTS, "Lato-Italic.ttf")))
    registerFontFamily("Lato", normal="Lato", bold="Lato-B", italic="Lato-I", boldItalic="Lato-B")

    MARGIN = 2.2 * cm
    avail = A4[0] - 2 * MARGIN

    body = ParagraphStyle("body", fontName="Lato", fontSize=10.5, leading=15.5,
                          textColor=HexColor(NEGRO), spaceAfter=7, alignment=4)
    h2s = ParagraphStyle("h2", fontName="Merri-B", fontSize=15, leading=18,
                         textColor=HexColor(AZUL), spaceBefore=16, spaceAfter=1)
    h3s = ParagraphStyle("h3", fontName="Merri-B", fontSize=12, leading=15,
                         textColor=HexColor(AZUL), spaceBefore=11, spaceAfter=3)
    li = ParagraphStyle("li", parent=body, leftIndent=14, bulletIndent=2, spaceAfter=3)
    cellh = ParagraphStyle("cellh", fontName="Lato-B", fontSize=9, leading=12, textColor=white)
    cell = ParagraphStyle("cell", fontName="Lato", fontSize=9, leading=12, textColor=HexColor(NEGRO))
    codest = ParagraphStyle("code", fontName="Courier", fontSize=8.6, leading=11.5, textColor=HexColor(NEGRO))
    quotest = ParagraphStyle("quote", fontName="Lato-I", fontSize=9.5, leading=14, textColor=HexColor(GRIS))
    tit = ParagraphStyle("tit", fontName="Merri-B", fontSize=25, leading=30,
                         textColor=HexColor(AZUL), alignment=1)
    metast = ParagraphStyle("meta", fontName="Lato", fontSize=11.5, leading=18,
                            textColor=HexColor(NEGRO), alignment=1)
    tochead = ParagraphStyle("tochead", fontName="Merri-B", fontSize=17, leading=21,
                             textColor=HexColor(AZUL), spaceAfter=2)
    toc0 = ParagraphStyle("toc0", fontName="Lato-B", fontSize=11, leading=20,
                          textColor=HexColor(AZUL), spaceBefore=4)
    toc1 = ParagraphStyle("toc1", fontName="Lato", fontSize=10, leading=16,
                          textColor=HexColor(NEGRO), leftIndent=16)

    class _Doc(SimpleDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                key = getattr(flowable, "_bk_key", None)
                if key and flowable.style.name in ("h2", "h3"):
                    lvl = 0 if flowable.style.name == "h2" else 1
                    txt = flowable.getPlainText()
                    self.notify("TOCEntry", (lvl, txt, self.page, key))
                    self.canv.bookmarkPage(key)
                    self.canv.addOutlineEntry(txt, key, level=lvl, closed=False)

    def il(t):
        out, pos = [], 0
        for m in INLINE.finditer(t):
            if m.start() > pos:
                out.append(html.escape(t[pos:m.start()], quote=False))
            tok = m.group(0)
            if tok.startswith("`"):
                out.append('<font face="Courier" size="9" backColor="%s"> %s </font>'
                           % (ZEBRA, html.escape(tok[1:-1], quote=False)))
            elif tok.startswith("**"):
                out.append("<b>%s</b>" % html.escape(tok[2:-2], quote=False))
            else:
                mm = LINKRE.match(tok)
                out.append('<link href="%s"><font color="%s"><u>%s</u></font></link>'
                           % (html.escape(mm.group(2), quote=False), AZUL,
                              html.escape(mm.group(1), quote=False)))
            pos = m.end()
        if pos < len(t):
            out.append(html.escape(t[pos:], quote=False))
        return "".join(out)

    def colw(rows):
        nc = max(len(r) for r in rows)
        mx = [1] * nc
        for r in rows:
            for j in range(nc):
                mx[j] = max(mx[j], len(r[j]) if j < len(r) else 0)
        tot = sum(mx)
        w = [max(avail * 0.10, avail * m / tot) for m in mx]
        f = avail / sum(w)
        return [x * f for x in w]

    story = []
    # portada
    story.append(Spacer(1, 3.2 * cm))
    img = Image(LOGO, width=9 * cm, height=9 * cm * 393 / 1600)
    img.hAlign = "CENTER"
    story.append(img)
    story.append(Spacer(1, 2.4 * cm))
    story.append(Paragraph(html.escape(h1 or "Informe", quote=False), tit))
    story.append(Spacer(1, 0.35 * cm))
    story.append(HRFlowable(width="32%", thickness=2.4, color=HexColor(DORADO),
                            spaceBefore=0, spaceAfter=14, hAlign="CENTER"))
    story.append(Spacer(1, 1.4 * cm))
    for k in ("Asignatura", "Autor", "Fecha"):
        if k in meta:
            story.append(Paragraph("<b>%s:</b> %s" % (k, html.escape(meta[k], quote=False)), metast))
    story.append(PageBreak())

    # indice
    toc = TableOfContents()
    toc.levelStyles = [toc0, toc1]
    story.append(Paragraph("Contenidos", tochead))
    story.append(HRFlowable(width="100%", thickness=1.4, color=HexColor(DORADO),
                            spaceBefore=1, spaceAfter=12))
    story.append(toc)
    story.append(PageBreak())

    # cuerpo
    hcount = 0
    for kind, val in blocks:
        if kind in ("h2", "h3"):
            hcount += 1
            key = "h%d" % hcount
            para = Paragraph(il(val), h2s if kind == "h2" else h3s)
            para._bk_key = key
            story.append(para)
            if kind == "h2":
                story.append(HRFlowable(width="100%", thickness=1.2, color=HexColor(DORADO),
                                        spaceBefore=1, spaceAfter=8))
        elif kind == "p":
            story.append(Paragraph(il(val), body))
        elif kind == "ul":
            for it in val:
                story.append(Paragraph(il(it), li, bulletText="•"))
        elif kind == "hr":
            story.append(Spacer(1, 4))
        elif kind == "quote":
            inner = Paragraph(il(val), quotest)
            t = Table([[inner]], colWidths=[avail])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HexColor("#FBFBFB")),
                ("LINEBEFORE", (0, 0), (0, -1), 3, HexColor(DORADO)),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]))
            story.append(t)
            story.append(Spacer(1, 7))
        elif kind == "code":
            pre = Preformatted(val, codest)
            t = Table([[pre]], colWidths=[avail])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HexColor(CODEBG)),
                ("BOX", (0, 0), (-1, -1), 0.5, HexColor("#E0E0E0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))
        elif kind == "table":
            rows = val
            data = []
            for ri, r in enumerate(rows):
                st = cellh if ri == 0 else cell
                data.append([Paragraph(il(c), st) for c in r])
            t = Table(data, colWidths=colw(rows), repeatRows=1)
            ts = [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor(AZUL)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LINEBELOW", (0, 0), (-1, -1), 0.4, HexColor("#DDDDDD")),
                ("LINEAFTER", (0, 0), (-2, -1), 0.4, HexColor("#EAEAEA")),
            ]
            for ri in range(1, len(rows)):
                if ri % 2 == 0:
                    ts.append(("BACKGROUND", (0, ri), (-1, ri), HexColor(ZEBRA)))
            t.setStyle(TableStyle(ts))
            story.append(t)
            story.append(Spacer(1, 10))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(HexColor(DORADO))
        canvas.setLineWidth(0.7)
        canvas.line(MARGIN, 1.45 * cm, A4[0] - MARGIN, 1.45 * cm)
        canvas.setFont("Lato", 8)
        canvas.setFillColor(HexColor(GRIS))
        canvas.drawString(MARGIN, 1.0 * cm, "Innovatech Solutions  ·  DSY1106")
        canvas.drawRightString(A4[0] - MARGIN, 1.0 * cm, "Página %d" % doc.page)
        canvas.restoreState()

    doc = _Doc(PDF, pagesize=A4, leftMargin=MARGIN, rightMargin=MARGIN,
               topMargin=2.0 * cm, bottomMargin=2.0 * cm,
               title="Informe tecnico - Innovatech Solutions",
               author=meta.get("Autor", ""))
    doc.multiBuild(story, onFirstPage=lambda c, d: None, onLaterPages=footer)
    print("PDF  ->", PDF)


# ---------------------------------------------------------------- DOCX
def build_docx(h1, meta, blocks):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    AZ = RGBColor(0x00, 0x21, 0x38)
    GR = RGBColor(0x66, 0x66, 0x66)
    NE = RGBColor(0x1A, 0x1A, 0x1A)
    WH = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    avail_cm = (sec.page_width - sec.left_margin - sec.right_margin)

    normal = doc.styles["Normal"]
    normal.font.name = "Lato"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NE
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), "Lato")

    def shade(el, fill):
        pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        pr.append(shd)

    def bottom_border(p, color):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "10")
        b.set(qn("w:space"), "3")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        pPr.append(pBdr)

    def hyperlink(p, url, text):
        rid = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
        h = OxmlElement("w:hyperlink")
        h.set(qn("r:id"), rid)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "002138")
        rPr.append(col)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "Lato")
        rf.set(qn("w:hAnsi"), "Lato")
        rPr.append(rf)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        h.append(r)
        p._p.append(h)

    def add_runs(p, text):
        pos = 0
        for m in INLINE.finditer(text):
            if m.start() > pos:
                p.add_run(text[pos:m.start()])
            tok = m.group(0)
            if tok.startswith("`"):
                r = p.add_run(tok[1:-1])
                r.font.name = "Courier New"
                r.font.size = Pt(9.5)
                rPr = r._element.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "EEEEEE")
                rPr.append(shd)
            elif tok.startswith("**"):
                p.add_run(tok[2:-2]).bold = True
            else:
                mm = LINKRE.match(tok)
                hyperlink(p, mm.group(2), mm.group(1))
            pos = m.end()
        if pos < len(text):
            p.add_run(text[pos:])

    # portada
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(5):
        doc.add_paragraph()
    pic2 = doc.add_paragraph()
    pic2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic2.add_run()
    run.add_picture(LOGO, width=Cm(9))
    for _ in range(2):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(h1 or "Informe")
    tr.font.name = "Merriweather"
    tr.font.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = AZ
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_border(rule, "FFB800")
    for _ in range(2):
        doc.add_paragraph()
    for k in ("Asignatura", "Autor", "Fecha"):
        if k in meta:
            mp = doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            b = mp.add_run(k + ": ")
            b.bold = True
            b.font.size = Pt(11.5)
            v = mp.add_run(meta[k])
            v.font.size = Pt(11.5)
    doc.add_page_break()

    # indice (campo TOC que Word actualiza al abrir)
    th = doc.add_paragraph()
    thr = th.add_run("Contenidos")
    thr.font.name = "Merriweather"
    thr.font.bold = True
    thr.font.size = Pt(17)
    thr.font.color.rgb = AZ
    bottom_border(th, "FFB800")
    th.paragraph_format.space_after = Pt(8)
    tp = doc.add_paragraph()
    trun = tp.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    fb.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fsep = OxmlElement("w:fldChar")
    fsep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t")
    ph.text = "Actualiza el índice: clic derecho sobre él y elige «Actualizar campos»."
    fsep.append(ph)
    fend = OxmlElement("w:fldChar")
    fend.set(qn("w:fldCharType"), "end")
    rel = trun._r
    rel.append(fb)
    rel.append(instr)
    rel.append(fsep)
    rel.append(fend)
    doc.add_page_break()

    # cuerpo
    for kind, val in blocks:
        if kind in ("h2", "h3"):
            p = doc.add_paragraph(style="Heading 1" if kind == "h2" else "Heading 2")
            r = p.add_run(re.sub(r"\*\*|`", "", val))
            r.font.name = "Merriweather"
            r.font.bold = True
            r.font.size = Pt(15 if kind == "h2" else 12.5)
            r.font.color.rgb = AZ
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(3)
            if kind == "h2":
                bottom_border(p, "FFB800")
        elif kind == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            add_runs(p, val)
        elif kind == "ul":
            for it in val:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, it)
        elif kind == "hr":
            pass
        elif kind == "quote":
            p = doc.add_paragraph()
            shade(p._p, "FBFBFB")
            bottom_border(p, "FFB800")
            p.paragraph_format.left_indent = Cm(0.3)
            r = p.add_run(re.sub(r"`", "", val))
            r.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = GR
        elif kind == "code":
            p = doc.add_paragraph()
            shade(p._p, "F5F5F5")
            p.paragraph_format.left_indent = Cm(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            clines = val.split("\n")
            for idx, cl in enumerate(clines):
                r = p.add_run(cl)
                r.font.name = "Courier New"
                r.font.size = Pt(9)
                if idx < len(clines) - 1:
                    r.add_break()
        elif kind == "table":
            rows = val
            nc = max(len(r) for r in rows)
            tb = doc.add_table(rows=len(rows), cols=nc)
            tb.style = "Table Grid"
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            tb.autofit = True
            for ri, r in enumerate(rows):
                for ci in range(nc):
                    cellp = tb.cell(ri, ci).paragraphs[0]
                    txt = r[ci] if ci < len(r) else ""
                    if ri == 0:
                        shade(tb.cell(ri, ci)._tc, "002138")
                        run = cellp.add_run(re.sub(r"\*\*|`", "", txt))
                        run.bold = True
                        run.font.color.rgb = WH
                        run.font.size = Pt(9.5)
                    else:
                        if ri % 2 == 0:
                            shade(tb.cell(ri, ci)._tc, "F4F4F4")
                        add_runs(cellp, txt)
                        for rr in cellp.runs:
                            rr.font.size = Pt(9.5)

    # footer
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Innovatech Solutions  ·  DSY1106  ·  Informe técnico")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GR

    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)

    doc.save(DOCX)
    print("DOCX ->", DOCX)


if __name__ == "__main__":
    with open(MD, encoding="utf-8") as f:
        h1, meta, blocks = parse_md(f.read())
    build_pdf(h1, meta, blocks)
    build_docx(h1, meta, blocks)
